import json
from pathlib import Path

import fitz
import pytest
from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.management import call_command
from django.urls import reverse

from apps.knowledge_base.chunking import ChunkingConfig, DocumentChunker
from apps.knowledge_base.exceptions import (
    DuplicateDocumentError,
    ExtractionError,
    UnapprovedDocumentError,
    UnsupportedDocumentError,
)
from apps.knowledge_base.extraction import (
    ExtractedBlock,
    ExtractionResult,
    ExtractorRegistry,
    PDFExtractor,
)
from apps.knowledge_base.ingestion import process_document
from apps.knowledge_base.metadata import (
    import_metadata,
    load_metadata,
    validate_metadata,
)
from apps.knowledge_base.models import (
    DocumentChunk,
    DocumentVersion,
    IngestionJob,
    KnowledgeDocument,
)
from apps.knowledge_base.registration import register_document_file
from apps.knowledge_base.validation import (
    sha256_file,
    validate_document_file,
)


@pytest.fixture
def reviewer():
    return get_user_model().objects.create_user(username="technical-reviewer")


@pytest.fixture
def metadata_payload() -> dict[str, str]:
    return {
        "document_id": "TEST-SAFETY-001",
        "title": "Test Robot Safety Procedure",
        "document_code": "TEST-001",
        "manufacturer": "Example Manufacturer",
        "document_type": "safety_handbook",
        "equipment_family": "Test Robot",
        "equipment_model": "TR-1",
        "subsystem": "safety",
        "edition": "1",
        "revision_date": "2026-07",
        "language": "English",
        "approval_status": "approved_for_prototype",
        "current_status": "verified_current",
        "safety_priority": "critical",
        "access_level": "internal",
        "notes": "Synthetic test fixture.",
    }


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "safe_fixture.pdf"
    document = fitz.open()
    page = document.new_page()
    text = """
1 ROBOT SAFETY
1.1 EMERGENCY STOP PROCEDURE
WARNING
Unexpected robot movement can cause serious injury.
1. Press the emergency stop before entering the safeguarded area.
2. Verify that robot motion has stopped.
3. Follow the approved lockout and tagout procedure.
CAUTION
Do not damage the emergency stop button during inspection.
NOTE
Record the inspection result in the approved checklist.
""".strip()
    page.insert_textbox(
        fitz.Rect(50, 50, 550, 760),
        text,
        fontsize=11,
    )
    document.set_toc(
        [
            [1, "1 ROBOT SAFETY", 1],
            [2, "1.1 EMERGENCY STOP PROCEDURE", 1],
        ]
    )
    document.save(path)
    document.close()
    return path


@pytest.mark.django_db
def test_metadata_import_requires_human_approval(
    metadata_payload: dict[str, str],
) -> None:
    document, created = import_metadata(metadata_payload)

    assert created is True
    assert document.document_id == "TEST-SAFETY-001"
    assert document.approval_status == KnowledgeDocument.ApprovalStatus.PENDING
    assert document.lifecycle_status == KnowledgeDocument.LifecycleStatus.UNDER_REVIEW
    assert (
        document.current_version_verification_status
        == KnowledgeDocument.VerificationStatus.VERIFIED_CURRENT
    )
    assert document.revision_label == "2026-07"
    assert document.revision_or_effective_date.isoformat() == "2026-07-01"


def test_metadata_yaml_uses_safe_mapping(
    tmp_path: Path,
    metadata_payload: dict[str, str],
) -> None:
    path = tmp_path / "metadata.yaml"
    path.write_text(
        "\n".join(
            f"{key}: {json.dumps(value)}" for key, value in metadata_payload.items()
        ),
        encoding="utf-8",
    )

    loaded = load_metadata(path)

    assert loaded["document_id"] == "TEST-SAFETY-001"


@pytest.mark.parametrize("extension", [".txt", ".md", ".markdown"])
def test_supported_text_formats_validate(tmp_path: Path, extension: str) -> None:
    path = tmp_path / f"fixture{extension}"
    path.write_text("# Safety\n\nApproved text.", encoding="utf-8")

    result = validate_document_file(path)

    assert result.extension == extension
    assert len(result.checksum) == 64


def test_pdf_validation_and_checksum_are_stable(sample_pdf: Path) -> None:
    first = validate_document_file(sample_pdf)
    second = sha256_file(sample_pdf)

    assert first.media_type == "application/pdf"
    assert first.checksum == second


def test_spoofed_pdf_is_rejected(tmp_path: Path) -> None:
    path = tmp_path / "spoofed.pdf"
    path.write_text("not a PDF", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentError):
        validate_document_file(path)


def test_docx_validation_and_extraction(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "fixture.docx"
    document = Document()
    document.add_heading("Robot Safety", level=1)
    document.add_paragraph("WARNING: Follow the approved procedure.")
    document.add_table(rows=1, cols=2).rows[0].cells[0].text = "Check"
    document.save(path)

    validated = validate_document_file(path)
    extraction = ExtractorRegistry().for_path(path).extract(path)

    assert validated.extension == ".docx"
    assert any(block.kind == "heading" for block in extraction.blocks)
    assert any(block.kind == "warning" for block in extraction.blocks)
    assert any(block.kind == "table" for block in extraction.blocks)


def test_pdf_extraction_preserves_safety_structure(sample_pdf: Path) -> None:
    extraction = PDFExtractor().extract(sample_pdf)
    kinds = [block.kind for block in extraction.blocks]
    text = extraction.text

    assert "heading" in kinds
    assert "warning" in kinds
    assert "caution" in kinds
    assert "procedure" in kinds
    assert "1.1 EMERGENCY STOP PROCEDURE" in text
    assert "Unexpected robot movement" in text
    assert "1. Press the emergency stop" in text


def test_chunking_preserves_warning_and_numbered_procedure(
    sample_pdf: Path,
) -> None:
    extraction = PDFExtractor().extract(sample_pdf)
    chunks = DocumentChunker(
        ChunkingConfig(
            target_tokens=400,
            overlap_tokens=50,
            minimum_tokens=20,
            maximum_tokens=500,
        )
    ).chunk(extraction, version_id="TEST-V1")

    combined = "\n".join(chunk.text for chunk in chunks)
    warning_chunk = next(chunk for chunk in chunks if chunk.contains_warning)
    assert "WARNING" in warning_chunk.text
    assert "Unexpected robot movement" in warning_chunk.text
    assert "1. Press the emergency stop" in warning_chunk.text
    assert "1. Press the emergency stop" in combined
    assert "2. Verify that robot motion has stopped" in combined
    assert warning_chunk.chapter == "1 ROBOT SAFETY"
    assert warning_chunk.section == "1.1 EMERGENCY STOP PROCEDURE"
    assert warning_chunk.page_start == 1


def test_near_identical_chunks_are_flagged() -> None:
    shared = " ".join(f"safety-term-{index}" for index in range(150))
    extraction = ExtractionResult(
        title="Duplicate fixture",
        page_count=2,
        blocks=[
            ExtractedBlock("heading", "1 FIRST", 1, "1 FIRST", "", 1),
            ExtractedBlock("paragraph", f"{shared} unique-a", 1, "1 FIRST"),
            ExtractedBlock("heading", "2 SECOND", 2, "2 SECOND", "", 1),
            ExtractedBlock("paragraph", f"{shared} unique-b", 2, "2 SECOND"),
        ],
    )

    chunks = DocumentChunker(ChunkingConfig(400, 50, 20, 500)).chunk(
        extraction, version_id="DUP-V1"
    )

    assert len(chunks) == 2
    assert chunks[1].duplicate_sequence == 1


@pytest.mark.django_db
def test_duplicate_document_checksum_is_detected(
    settings,
    tmp_path: Path,
    sample_pdf: Path,
    metadata_payload: dict[str, str],
) -> None:
    settings.MEDIA_ROOT = tmp_path / "media"
    first, _ = import_metadata(metadata_payload)
    register_document_file(first, sample_pdf)
    second_payload = {**metadata_payload, "document_id": "TEST-SAFETY-002"}
    second, _ = import_metadata(second_payload)

    with pytest.raises(DuplicateDocumentError):
        register_document_file(second, sample_pdf)


@pytest.mark.django_db
def test_unapproved_document_cannot_be_processed(
    settings,
    tmp_path: Path,
    sample_pdf: Path,
    metadata_payload: dict[str, str],
) -> None:
    settings.MEDIA_ROOT = tmp_path / "media"
    document, _ = import_metadata(metadata_payload)
    register_document_file(document, sample_pdf)

    with pytest.raises(UnapprovedDocumentError):
        process_document(document.pk)

    assert IngestionJob.objects.count() == 0


@pytest.mark.django_db
def test_processing_creates_inherited_stable_chunks(
    settings,
    tmp_path: Path,
    sample_pdf: Path,
    metadata_payload: dict[str, str],
    reviewer,
) -> None:
    settings.MEDIA_ROOT = tmp_path / "media"
    document, _ = import_metadata(metadata_payload)
    register_document_file(document, sample_pdf)
    document.approval_status = KnowledgeDocument.ApprovalStatus.APPROVED
    document.lifecycle_status = KnowledgeDocument.LifecycleStatus.APPROVED
    document.approved_by = reviewer
    document.save()

    first_job = process_document(document.pk)
    first_ids = list(
        DocumentChunk.objects.filter(document=document).values_list(
            "chunk_id",
            flat=True,
        )
    )
    second_job = process_document(document.pk, reprocess=True)
    second_ids = list(
        DocumentChunk.objects.filter(document=document).values_list(
            "chunk_id",
            flat=True,
        )
    )

    assert first_job.status == IngestionJob.Status.COMPLETED
    assert second_job.status == IngestionJob.Status.COMPLETED
    assert first_ids == second_ids
    chunk = DocumentChunk.objects.get(document=document, contains_warning=True)
    assert chunk.manufacturer == document.manufacturer
    assert chunk.equipment_family == document.equipment_family
    assert chunk.subsystem == document.subsystem
    assert chunk.safety_priority == KnowledgeDocument.SafetyPriority.CRITICAL
    assert chunk.page_start == 1
    assert len(chunk.content_hash) == 64


@pytest.mark.django_db
def test_failed_extraction_is_recorded(
    settings,
    tmp_path: Path,
    metadata_payload: dict[str, str],
    reviewer,
) -> None:
    settings.MEDIA_ROOT = tmp_path / "media"
    document, _ = import_metadata(metadata_payload)
    document.approval_status = KnowledgeDocument.ApprovalStatus.APPROVED
    document.lifecycle_status = KnowledgeDocument.LifecycleStatus.APPROVED
    document.approved_by = reviewer
    document.save()
    version = DocumentVersion(
        version_id="BROKEN-V1",
        document=document,
        source_filename="broken.pdf",
        checksum="a" * 64,
        file_size=12,
        media_type="application/pdf",
    )
    version.source_file.save(
        "broken.pdf",
        ContentFile(b"%PDF-broken"),
        save=True,
    )

    with pytest.raises(ExtractionError):
        process_document(document.pk)

    job = IngestionJob.objects.get()
    document.refresh_from_db()
    assert job.status == IngestionJob.Status.FAILED
    assert job.errors
    assert document.processing_status == KnowledgeDocument.ProcessingStatus.FAILED


@pytest.mark.django_db
def test_validate_metadata_flags_approved_document_without_reviewer(
    metadata_payload: dict[str, str],
) -> None:
    document, _ = import_metadata(metadata_payload)
    document.approval_status = KnowledgeDocument.ApprovalStatus.APPROVED
    document.save()

    errors = validate_metadata(document)

    assert "Approved documents must identify approved_by." in errors


@pytest.mark.django_db
def test_export_chunks_command(
    settings,
    tmp_path: Path,
    sample_pdf: Path,
    metadata_payload: dict[str, str],
    reviewer,
) -> None:
    settings.MEDIA_ROOT = tmp_path / "media"
    document, _ = import_metadata(metadata_payload)
    register_document_file(document, sample_pdf)
    document.approval_status = KnowledgeDocument.ApprovalStatus.APPROVED
    document.lifecycle_status = KnowledgeDocument.LifecycleStatus.APPROVED
    document.approved_by = reviewer
    document.save()
    process_document(document.pk)
    output = tmp_path / "chunks.json"

    call_command(
        "export_document_chunks",
        document.pk,
        "--output",
        str(output),
    )

    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload["document_id"] == document.pk
    assert payload["chunk_count"] == DocumentChunk.objects.count()


@pytest.mark.django_db
def test_admin_preview_is_staff_protected(
    client,
    metadata_payload: dict[str, str],
) -> None:
    document, _ = import_metadata(metadata_payload)
    url = reverse(
        "admin:knowledge_base_knowledgedocument_preview",
        args=[document.pk],
    )

    anonymous_response = client.get(url)
    assert anonymous_response.status_code == 302

    ordinary_user = get_user_model().objects.create_user(username="ordinary")
    client.force_login(ordinary_user)
    ordinary_response = client.get(url)
    assert ordinary_response.status_code == 302

    staff_user = get_user_model().objects.create_user(
        username="staff",
        is_staff=True,
        is_superuser=True,
    )
    client.force_login(staff_user)
    staff_response = client.get(url)
    assert staff_response.status_code == 200
    assert document.document_id.encode() in staff_response.content
