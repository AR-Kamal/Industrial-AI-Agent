"""Provider-neutral document extraction with local format adapters."""

import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from statistics import median
from typing import Protocol, TypedDict

import fitz
from docx import Document

from .exceptions import ExtractionError, ManualReviewRequired

HEADING_PATTERN = re.compile(r"^\d+(?:\.\d+)*\s+[A-Z][A-Za-z0-9 ,/&()'-]+$")
PROCEDURE_PATTERN = re.compile(r"^(?:\(?\d+\)?[.)]\s+|[a-zA-Z][.)]\s+|[•●▪]\s+|z\s+)")
FIGURE_PATTERN = re.compile(r"^(?:figure|fig\.|table)\s+\d+", re.IGNORECASE)
PAGE_NUMBER_PATTERN = re.compile(r"^(?:page\s+)?[ivxlcdm\d]+$", re.IGNORECASE)


class PageLine(TypedDict):
    text: str
    size: float
    font: str
    top_ratio: float
    bottom_ratio: float


@dataclass(frozen=True)
class ProcessingWarning:
    code: str
    message: str
    page: int | None = None

    def as_dict(self) -> dict[str, str | int | None]:
        return {"code": self.code, "message": self.message, "page": self.page}


@dataclass
class ExtractedBlock:
    kind: str
    text: str
    page_number: int | None = None
    chapter: str = ""
    section: str = ""
    heading_level: int = 0


@dataclass
class ExtractionResult:
    title: str
    blocks: list[ExtractedBlock]
    page_count: int
    warnings: list[ProcessingWarning] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        parts: list[str] = []
        current_page: int | None = None
        for block in self.blocks:
            if block.page_number != current_page and block.page_number is not None:
                current_page = block.page_number
                parts.append(f"\n[Page {current_page}]")
            parts.append(block.text)
        return "\n\n".join(parts).strip()


class DocumentExtractor(Protocol):
    def extract(self, path: Path) -> ExtractionResult:
        """Extract structured blocks without embedding or indexing."""

        ...


class ExtractorRegistry:
    def __init__(self) -> None:
        self._extractors: dict[str, DocumentExtractor] = {
            ".pdf": PDFExtractor(),
            ".docx": DocxExtractor(),
            ".txt": PlainTextExtractor(markdown=False),
            ".md": PlainTextExtractor(markdown=True),
            ".markdown": PlainTextExtractor(markdown=True),
        }

    def for_path(self, path: Path) -> DocumentExtractor:
        try:
            return self._extractors[path.suffix.lower()]
        except KeyError as exc:
            raise ExtractionError(
                "No extractor is registered for this format."
            ) from exc


class PDFExtractor:
    def extract(self, path: Path) -> ExtractionResult:
        try:
            document = fitz.open(path)
        except (fitz.FileDataError, RuntimeError) as exc:
            raise ExtractionError("The PDF could not be opened.") from exc

        try:
            raw_pages = [self._page_lines(page) for page in document]
            repeated_margins = self._repeated_margin_text(
                raw_pages, document.page_count
            )
            outline = self._outline(document)
            warnings: list[ProcessingWarning] = []
            blocks: list[ExtractedBlock] = []
            empty_pages = 0
            current_chapter = ""
            current_section = ""

            for page_index, (page, lines) in enumerate(
                zip(document, raw_pages, strict=False), start=1
            ):
                usable_lines = [
                    line
                    for line in lines
                    if self._margin_key(line) not in repeated_margins
                    and not PAGE_NUMBER_PATTERN.fullmatch(line["text"].strip())
                ]
                page_text_length = sum(
                    len(line["text"].strip()) for line in usable_lines
                )
                if page_text_length == 0:
                    empty_pages += 1
                    warnings.append(
                        ProcessingWarning(
                            "empty_page",
                            "The page contains no extractable text.",
                            page_index,
                        )
                    )
                    if page.get_images(full=True):
                        warnings.append(
                            ProcessingWarning(
                                "possible_scanned_page",
                                "Page images have no usable text; OCR was not run.",
                                page_index,
                            )
                        )

                image_count = len(page.get_images(full=True))
                drawing_count = len(page.get_drawings())
                if image_count or drawing_count >= 40:
                    warnings.append(
                        ProcessingWarning(
                            "diagram_present",
                            "Images or dense vector drawings require manual review.",
                            page_index,
                        )
                    )

                (
                    page_blocks,
                    current_chapter,
                    current_section,
                ) = self._classify_lines(
                    usable_lines,
                    page_index,
                    outline,
                    current_chapter,
                    current_section,
                )
                blocks.extend(page_blocks)
                blocks.extend(
                    self._extract_tables(
                        page,
                        page_index,
                        warnings,
                        chapter=current_chapter,
                        section=current_section,
                    )
                )

            if not blocks or sum(len(block.text) for block in blocks) < 100:
                raise ManualReviewRequired(
                    "The PDF contains no usable text. OCR is outside this milestone."
                )
            if empty_pages > max(3, document.page_count // 2):
                raise ManualReviewRequired(
                    "Too many PDF pages have no usable text; manual review is required."
                )

            title = str(document.metadata.get("title") or "").strip()
            if not title:
                title = next(
                    (
                        block.text
                        for block in blocks
                        if block.kind in {"title", "heading"}
                    ),
                    path.stem,
                )
            return clean_extraction(
                ExtractionResult(
                    title=title,
                    blocks=blocks,
                    page_count=document.page_count,
                    warnings=warnings,
                )
            )
        finally:
            document.close()

    @staticmethod
    def _page_lines(page: fitz.Page) -> list[PageLine]:
        result: list[PageLine] = []
        page_dict = page.get_text("dict", sort=True)
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = "".join(str(span.get("text", "")) for span in spans).strip()
                if not text:
                    continue
                sizes = [float(span.get("size", 0)) for span in spans]
                fonts = " ".join(str(span.get("font", "")) for span in spans)
                bbox = line.get("bbox", (0, 0, 0, 0))
                result.append(
                    {
                        "text": text,
                        "size": max(sizes, default=0),
                        "font": fonts,
                        "top_ratio": float(bbox[1]) / max(page.rect.height, 1),
                        "bottom_ratio": float(bbox[3]) / max(page.rect.height, 1),
                    }
                )
        return result

    @staticmethod
    def _margin_key(line: PageLine) -> str:
        text = re.sub(r"\s+", " ", str(line["text"])).strip().casefold()
        if float(line["top_ratio"]) <= 0.12 or float(line["bottom_ratio"]) >= 0.88:
            return text
        return ""

    def _repeated_margin_text(
        self,
        pages: list[list[PageLine]],
        page_count: int,
    ) -> set[str]:
        candidates: Counter[str] = Counter()
        for lines in pages:
            candidates.update(
                {
                    key
                    for line in lines
                    if (key := self._margin_key(line)) and len(key) > 2
                }
            )
        threshold = max(3, int(page_count * 0.25))
        return {text for text, count in candidates.items() if count >= threshold}

    @staticmethod
    def _outline(document: fitz.Document) -> dict[int, list[tuple[int, str]]]:
        outline: dict[int, list[tuple[int, str]]] = {}
        for level, title, page in document.get_toc():
            clean_title = re.sub(r"\s+", " ", title.replace("\r", " ")).strip()
            outline.setdefault(int(page), []).append((int(level), clean_title))
        return outline

    def _classify_lines(
        self,
        lines: list[PageLine],
        page_number: int,
        outline: dict[int, list[tuple[int, str]]],
        initial_chapter: str,
        initial_section: str,
    ) -> tuple[list[ExtractedBlock], str, str]:
        if not lines:
            return [], initial_chapter, initial_section
        body_size = median(float(line["size"]) for line in lines)
        page_headings = outline.get(page_number, [])
        chapter = initial_chapter
        section = initial_section
        blocks: list[ExtractedBlock] = []

        index = 0
        while index < len(lines):
            line = lines[index]
            text = normalize_line(str(line["text"]))
            if not text:
                index += 1
                continue
            matched_heading = self._match_outline_heading(text, page_headings)
            consumed = 1
            if matched_heading is None:
                maximum = min(4, len(lines) - index)
                for candidate_length in range(maximum, 1, -1):
                    combined = " ".join(
                        normalize_line(str(candidate["text"]))
                        for candidate in lines[index : index + candidate_length]
                    )
                    matched_heading = self._match_outline_heading(
                        combined,
                        page_headings,
                    )
                    if matched_heading is not None:
                        consumed = candidate_length
                        break
            if matched_heading:
                level, heading = matched_heading
                if level == 1:
                    chapter = heading
                    section = ""
                else:
                    section = heading
                if chapter == "TABLE OF CONTENTS":
                    index += consumed
                    continue
                blocks.append(
                    ExtractedBlock(
                        "heading",
                        heading,
                        page_number,
                        chapter,
                        section,
                        level,
                    )
                )
                index += consumed
                continue

            kind = classify_text(
                text,
                font_size=float(line["size"]),
                body_size=body_size,
                is_bold="bold" in str(line["font"]).casefold(),
            )
            if chapter == "TABLE OF CONTENTS" and kind not in {"warning", "caution"}:
                index += 1
                continue
            blocks.append(
                ExtractedBlock(
                    kind,
                    text,
                    page_number,
                    chapter,
                    section,
                )
            )
            index += 1
        return blocks, chapter, section

    @staticmethod
    def _match_outline_heading(
        text: str,
        headings: list[tuple[int, str]],
    ) -> tuple[int, str] | None:
        normalized = normalize_for_comparison(text)
        for level, heading in sorted(headings, reverse=True):
            heading_normalized = normalize_for_comparison(heading)
            if normalized == heading_normalized or (
                len(heading_normalized) >= 8 and heading_normalized in normalized
            ):
                return level, heading
        return None

    @staticmethod
    def _extract_tables(
        page: fitz.Page,
        page_number: int,
        warnings: list[ProcessingWarning],
        *,
        chapter: str,
        section: str,
    ) -> list[ExtractedBlock]:
        try:
            tables = page.find_tables().tables
        except (AttributeError, RuntimeError):
            return []
        blocks: list[ExtractedBlock] = []
        for table in tables:
            rows = table.extract()
            rendered_rows = [
                "| " + " | ".join((cell or "").strip() for cell in row) + " |"
                for row in rows
                if any((cell or "").strip() for cell in row)
            ]
            if rendered_rows:
                blocks.append(
                    ExtractedBlock(
                        "table",
                        "\n".join(rendered_rows),
                        page_number,
                        chapter,
                        section,
                    )
                )
                warnings.append(
                    ProcessingWarning(
                        "table_extraction_uncertain",
                        "A table was extracted and requires layout verification.",
                        page_number,
                    )
                )
        return blocks


class DocxExtractor:
    def extract(self, path: Path) -> ExtractionResult:
        try:
            document = Document(str(path))
        except (ValueError, OSError) as exc:
            raise ExtractionError("The DOCX could not be opened.") from exc

        blocks: list[ExtractedBlock] = []
        chapter = ""
        section = ""
        for paragraph in document.paragraphs:
            text = normalize_line(paragraph.text)
            if not text:
                continue
            style = paragraph.style.name.casefold() if paragraph.style else ""
            heading_match = re.search(r"heading\s*(\d+)", style)
            if heading_match:
                level = int(heading_match.group(1))
                if level == 1:
                    chapter = text
                    section = ""
                else:
                    section = text
                kind = "heading"
            else:
                level = 0
                kind = classify_text(text)
            blocks.append(ExtractedBlock(kind, text, None, chapter, section, level))

        for table in document.tables:
            rows = [
                "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
                for row in table.rows
            ]
            blocks.append(ExtractedBlock("table", "\n".join(rows)))

        if not blocks:
            raise ManualReviewRequired("The DOCX contains no usable text.")
        return clean_extraction(
            ExtractionResult(
                title=next(
                    (block.text for block in blocks if block.kind == "heading"),
                    path.stem,
                ),
                blocks=blocks,
                page_count=0,
                warnings=[
                    ProcessingWarning(
                        "missing_page_references",
                        "DOCX extraction does not provide reliable page numbers.",
                    )
                ],
            )
        )


class PlainTextExtractor:
    def __init__(self, *, markdown: bool) -> None:
        self.markdown = markdown

    def extract(self, path: Path) -> ExtractionResult:
        try:
            content = path.read_text(encoding="utf-8")
        except (OSError, UnicodeDecodeError) as exc:
            raise ExtractionError("The text document could not be read.") from exc

        blocks: list[ExtractedBlock] = []
        chapter = ""
        section = ""
        for raw_block in re.split(r"\n\s*\n", content):
            text = normalize_line(raw_block)
            if not text:
                continue
            markdown_heading = re.match(r"^(#{1,6})\s+(.+)$", text)
            if self.markdown and markdown_heading:
                level = len(markdown_heading.group(1))
                text = markdown_heading.group(2).strip()
                kind = "heading"
            else:
                level = 1 if HEADING_PATTERN.fullmatch(text) else 0
                kind = "heading" if level else classify_text(text)
            if kind == "heading":
                if level == 1:
                    chapter = text
                    section = ""
                else:
                    section = text
            blocks.append(ExtractedBlock(kind, text, None, chapter, section, level))
        if not blocks:
            raise ManualReviewRequired("The text document contains no usable text.")
        return clean_extraction(
            ExtractionResult(
                title=next(
                    (block.text for block in blocks if block.kind == "heading"),
                    path.stem,
                ),
                blocks=blocks,
                page_count=0,
                warnings=[
                    ProcessingWarning(
                        "missing_page_references",
                        "Plain-text extraction has no page references.",
                    )
                ],
            )
        )


def classify_text(
    text: str,
    *,
    font_size: float = 0,
    body_size: float = 0,
    is_bold: bool = False,
) -> str:
    upper = text.strip().upper()
    if upper in {"WARNING", "WARINIG"} or upper.startswith(
        ("WARNING ", "WARNING:", "WARINIG ", "WARINIG:")
    ):
        return "warning"
    if upper == "CAUTION" or upper.startswith(("CAUTION ", "CAUTION:")):
        return "caution"
    if upper == "NOTE" or upper.startswith(("NOTE ", "NOTE:")):
        return "note"
    if FIGURE_PATTERN.match(text):
        return "figure_caption"
    if PROCEDURE_PATTERN.match(text):
        return "procedure"
    if HEADING_PATTERN.fullmatch(text) or (
        body_size and font_size > body_size * 1.15 and (is_bold or len(text) < 100)
    ):
        return "heading"
    if text.startswith("|") and "|" in text[1:]:
        return "table"
    return "paragraph"


def normalize_line(text: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    if not lines:
        return ""
    result = lines[0]
    for line in lines[1:]:
        if result.endswith("-") and line[:1].islower():
            result = result[:-1] + line
        elif (
            not re.search(r"[.!?:;]$", result)
            and not PROCEDURE_PATTERN.match(line)
            and not result.startswith("|")
        ):
            result += " " + line
        else:
            result += "\n" + line
    return re.sub(r" {2,}", " ", result).strip()


def normalize_for_comparison(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", text.casefold()).strip()


def clean_extraction(result: ExtractionResult) -> ExtractionResult:
    cleaned: list[ExtractedBlock] = []
    seen: set[str] = set()
    pending_safety: ExtractedBlock | None = None

    for block in result.blocks:
        block.text = normalize_line(block.text)
        if not block.text:
            continue
        if block.kind in {"warning", "caution"} and len(block.text.split()) <= 3:
            pending_safety = block
            continue
        if pending_safety is not None:
            if block.kind != "heading":
                block = ExtractedBlock(
                    pending_safety.kind,
                    f"{pending_safety.text}\n{block.text}",
                    block.page_number,
                    block.chapter or pending_safety.chapter,
                    block.section or pending_safety.section,
                )
            else:
                cleaned.append(pending_safety)
            pending_safety = None

        comparison = normalize_for_comparison(block.text)
        protected = block.kind in {
            "warning",
            "caution",
            "procedure",
            "heading",
            "table",
        }
        if comparison in seen and not protected:
            continue
        if comparison:
            seen.add(comparison)
        cleaned.append(block)

    if pending_safety is not None:
        cleaned.append(pending_safety)

    result.blocks = merge_wrapped_paragraphs(cleaned)
    if not any(block.kind == "heading" for block in result.blocks):
        result.warnings.append(
            ProcessingWarning(
                "missing_headings",
                "No reliable headings were detected.",
            )
        )
    return result


def merge_wrapped_paragraphs(
    blocks: list[ExtractedBlock],
) -> list[ExtractedBlock]:
    merged: list[ExtractedBlock] = []
    for block in blocks:
        if (
            merged
            and block.kind == "paragraph"
            and merged[-1].kind == "paragraph"
            and block.page_number == merged[-1].page_number
            and block.section == merged[-1].section
            and len(merged[-1].text) < 1200
        ):
            merged[-1].text = f"{merged[-1].text} {block.text}"
        else:
            merged.append(block)
    return merged
